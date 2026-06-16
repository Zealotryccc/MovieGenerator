from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

path = os.path.join(os.path.dirname(__file__), 'MovieGenerator_docs.docx')
doc = Document(path)

# Find where the title page ends (first heading)
first_heading_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading'):
        first_heading_idx = i
        break

for i, p in enumerate(doc.paragraphs):
    # Skip title page
    if first_heading_idx is not None and i >= first_heading_idx:
        if p.style.name.startswith('Heading'):
            for run in p.runs:
                run.font.size = Pt(16)
        else:
            for run in p.runs:
                run.font.size = Pt(14)

doc.save(path)
print('Done')
